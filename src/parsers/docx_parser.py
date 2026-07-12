"""
Word 解析器 — 基于 python-docx

提取段落（含样式层级）+ 表格(Markdown)。
Heading 1/2/3 样式直接映射为 h1/h2/h3 标题层级。
"""

import io
from typing import NamedTuple


class DocxElement(NamedTuple):
    type: str          # "paragraph" | "table"
    text: str          # 段落的纯文本，或表格的 Markdown
    heading_level: int # 0=正文, 1=标题1, 2=标题2, ...
    is_bold: bool


class DocxParseResult(NamedTuple):
    elements: list[DocxElement]
    full_text: str
    total_chars: int
    title: str          # 文档标题（第一个 Heading 1 或文件名）
    error: str | None


def parse_docx(file_bytes: bytes) -> DocxParseResult:
    """解析 Word 文档字节流。"""
    from docx import Document

    elements: list[DocxElement] = []
    title = ""
    error = None

    try:
        doc = Document(io.BytesIO(file_bytes))

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            heading_level = _get_heading_level(para.style.name if para.style else "")
            is_bold = False
            if para.runs:
                is_bold = para.runs[0].bold or False

            if heading_level == 1 and not title:
                title = text

            elements.append(DocxElement(
                type="paragraph",
                text=text,
                heading_level=heading_level,
                is_bold=is_bold,
            ))

        for table in doc.tables:
            md = _docx_table_to_markdown(table)
            if md:
                elements.append(DocxElement(
                    type="table", text=md, heading_level=0, is_bold=False,
                ))

    except Exception as e:
        error = f"Word 解析失败: {e}"

    full_text = _elements_to_text(elements)
    return DocxParseResult(
        elements=elements,
        full_text=full_text,
        total_chars=len(full_text),
        title=title,
        error=error,
    )


def _get_heading_level(style_name: str) -> int:
    """从 Word 样式名提取标题层级。"""
    name = style_name.lower()
    # 英文样式
    for level in range(1, 7):
        if f"heading {level}" in name or f"heading{level}" in name:
            return level
    # 中文样式
    for level in range(1, 7):
        if f"标题 {level}" in name or f"标题{level}" in name:
            return level
    # 大纲级别样式
    for level in range(1, 7):
        if f"toc {level}" in name or f"toc{level}" in name:
            return level
    return 0


def _docx_table_to_markdown(table) -> str:
    """Word 表格 → Markdown 字符串。"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows or not rows[0]:
        return ""

    header = rows[0]
    sep = ["---"] * len(header)
    body = rows[1:]

    lines = [
        "| " + " | ".join(header) + " |",
        "|" + "|".join(sep) + "|",
    ]
    for row in body:
        padded = row + [""] * (len(header) - len(row))
        lines.append("| " + " | ".join(padded[:len(header)]) + " |")

    return "\n".join(lines)


def _elements_to_text(elements: list[DocxElement]) -> str:
    """将结构化元素列表拼成全文。"""
    parts: list[str] = []
    for el in elements:
        prefix = "#" * el.heading_level + " " if el.heading_level else ""
        if el.type == "table":
            parts.append(el.text)
        else:
            parts.append(prefix + el.text)
    return "\n\n".join(parts)
